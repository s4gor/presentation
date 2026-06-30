#!/usr/bin/env python3
"""
Main submission case paper - Team 14
Horizon_2035_Scenario_Case_Team14.docx
Plain language, full topics, embedded visuals.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_LINE_SPACING
from PIL import Image, ImageDraw, ImageFont

BASE = Path("/Users/emran/Desktop/pptx")
ASSETS = BASE / "team_assets"
OUTPUT = BASE / "Horizon_2035_Scenario_Case_Team14.docx"


def _font(size, bold=False):
    name = "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf"
    try:
        return ImageFont.truetype(name, size)
    except OSError:
        return ImageFont.load_default()


def make_matrix_image():
    ASSETS.mkdir(parents=True, exist_ok=True)
    path = ASSETS / "matrix_2x2.png"
    W, H = 1200, 900
    img = Image.new("RGB", (W, H), "#FAFAF8")
    draw = ImageDraw.Draw(img)
    tf, lf, cf = _font(28, True), _font(20, True), _font(17)
    sf, af = _font(14), _font(16)

    draw.text((40, 24), "2x2 Scenario Matrix - Team 14 - Horizon 2035", fill="#1a1a1a", font=tf)
    gx, gy, gw, gh = 280, 120, 860, 680
    cw, ch = gw // 2, gh // 2
    palette = {
        "gate": ("#FDE8E8", "#C44D6A", "3. Gatekeeper Web"),
        "cite": ("#E6F5F4", "#0A7C78", "1. Citation Economy"),
        "slow": ("#FEF3C7", "#B8943F", "4. Slow Adaptation"),
        "last": ("#ECFDF5", "#059669", "2. Last-Click Web"),
    }
    cells = [
        (0, 0, "gate", "AI-first search.\nFunnels still expected.\nUsers rarely click.\nPlatforms pick winners."),
        (1, 0, "cite", "AI-first search.\nZero-click wins.\nBe cited in answers.\nOwn specs and data."),
        (0, 1, "slow", "Google + AI coexist.\nZero-click creeps up.\nTeams adapt too late."),
        (1, 1, "last", "Google + AI coexist.\nWebsites still convert.\nSEO = trust + structure."),
    ]
    for col, row, key, body in cells:
        x, y = gx + col * cw, gy + row * ch
        bg, accent, title = palette[key]
        draw.rectangle([x, y, x + cw - 4, y + ch - 4], fill=bg, outline=accent, width=3)
        draw.text((x + 16, y + 14), title, fill=accent, font=lf)
        draw.multiline_text((x + 16, y + 48), body, fill="#333", font=cf, spacing=5)

    draw.text((gx + 80, 88), "Website funnels still dominate (left)", fill="#555", font=af)
    draw.text((gx + cw + 60, 88), "Zero-click / citations win (right)", fill="#555", font=af)
    draw.text((20, gy + 80), "AI-native\ndiscovery\n(top row)", fill="#555", font=af)
    draw.text((20, gy + ch + 80), "Hybrid\ndiscovery\n(bottom row)", fill="#555", font=af)
    draw.text((gx + 200, H - 40), "Horizontal axis = VALUE CAPTURE", fill="#888", font=sf)
    draw.text((8, gy + 300), "Vertical\n= DISCOVERY", fill="#888", font=sf)
    img.save(path)
    return path


def make_system_image():
    path = ASSETS / "system_picture.png"
    W, H = 1000, 700
    img = Image.new("RGB", (W, H), "#FAFAF8")
    draw = ImageDraw.Draw(img)
    tf, lf, bf = _font(26, True), _font(18, True), _font(15)
    draw.text((40, 20), "System Picture - Four Layers", fill="#1a1a1a", font=tf)

    layers = [
        ("General environment", "#E8EAF0", "GenAI, copyright, Gartner forecasts, energy, global platforms"),
        ("Specific environment", "#D6EEF0", "EU/Germany, SISTRIX, SparkToro EU, DMA, AI Act, industry split"),
        ("Value chain", "#C5E8E6", "Users > Platforms > Sites > Agencies > B2B funnels > Analytics"),
        ("Core (centre)", "#0A7C78", "Organic discoverability when AI mediates search"),
    ]
    y = 80
    margin = 40
    for i, (title, color, body) in enumerate(layers):
        inset = i * 35
        x1, x2 = margin + inset, W - margin - inset
        h = 120 if i < 3 else 100
        text_color = "white" if i == 3 else "#222"
        draw.rounded_rectangle([x1, y, x2, y + h], radius=12, fill=color)
        draw.text((x1 + 20, y + 16), title, fill=text_color, font=lf)
        draw.text((x1 + 20, y + 48), body, fill=text_color, font=bf)
        y += h + 14

    img.save(path)
    return path


def make_data_image():
    path = ASSETS / "key_statistics.png"
    W, H = 1000, 520
    img = Image.new("RGB", (W, H), "#FAFAF8")
    draw = ImageDraw.Draw(img)
    tf, lf, vf = _font(24, True), _font(14), _font(22, True)
    draw.text((40, 16), "Key Statistics (sources in paper)", fill="#1a1a1a", font=tf)

    stats = [
        ("59.7%", "EU zero-click", "#C44D6A"),
        ("374", "clicks / 1000 EU searches", "#0A7C78"),
        ("20%", "DE queries with AIO", "#2563EB"),
        ("27 to 11%", "CTR on #1 when AIO", "#B8943F"),
        ("-25%", "search volume (Gartner)", "#6D4CDB"),
        ("265M", "lost clicks/mo Germany", "#059669"),
    ]
    cols, pad, bw, bh = 3, 24, 300, 120
    for i, (num, label, color) in enumerate(stats):
        col, row = i % cols, i // cols
        x = 40 + col * (bw + pad)
        y = 70 + row * (bh + pad)
        draw.rounded_rectangle([x, y, x + bw, y + bh], radius=10, fill="#FFF", outline=color, width=2)
        draw.text((x + 20, y + 22), num, fill=color, font=vf)
        draw.text((x + 20, y + 62), label, fill="#444", font=lf)

    img.save(path)
    return path


def body(doc, text, space=8):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(space)
    return p


def h1(doc, t):
    return doc.add_heading(t, level=1)


def h2(doc, t):
    return doc.add_heading(t, level=2)


def bullet(doc, t):
    p = doc.add_paragraph(t, style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    return p


def fig(doc, path, width=6.0, caption=""):
    doc.add_picture(str(path), width=Inches(width))
    if caption:
        p = doc.add_paragraph(caption)
        p.paragraph_format.space_after = Pt(12)
        for r in p.runs:
            r.italic = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def build():
    matrix_img = make_matrix_image()
    system_img = make_system_image()
    data_img = make_data_image()

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1.1)
        s.right_margin = Inches(1.1)

    t = doc.add_heading("SEO in the Age of AI", 0)
    t.alignment = 1
    body(doc, "Scenario Planning Case Study | Horizon 2035 | Brief B")
    body(doc, "Team 14: Emran (coordinator), Arshi, Ashik, Ozlem, Wasay")
    body(doc, "Scenario Management | Thomas Chermack scenario planning cycle")
    doc.add_paragraph()

    # ABSTRACT
    h1(doc, "Abstract")
    body(doc,
        "For years, SEO meant ranking on Google, winning the click, and converting on your website. "
        "That model still works in places, but it is no longer the full picture. A large share of searches "
        "in the EU end without any click to the open web. AI Overviews sit above organic results. "
        "Tools like ChatGPT are becoming the first stop for work research.")
    body(doc,
        "Team 14 applied Chermack's scenario planning cycle to Brief B. We scanned the external environment (STEEP), "
        "separated Givens from Drivers, interviewed two practitioners outside our team, and built a 2x2 matrix "
        "with four scenarios set in 2035. This paper is not a prediction. It is a structured way to prepare "
        "for different plausible futures in organic discoverability when AI mediates search.")

    # 1 INTRO
    h1(doc, "1. Introduction")
    h2(doc, "1.1 Why this project matters")
    body(doc,
        "Brief B asks what happens to organic discoverability when AI sits between brands and users. "
        "The audience is marketing and communications professionals: brand managers, agency leads, "
        "publisher strategists, and B2B marketers who still defend budgets every quarter.")
    body(doc,
        "Universities often teach SEO as if traffic is the only KPI. Our data and interviews suggest "
        "visibility can mean citation inside an AI answer, not a visit to your site. That shift is "
        "strategic, not just technical.")
    h2(doc, "1.2 What we are not claiming")
    body(doc,
        "We do not claim Google disappears by 2030 or that SEO vanishes. We map four internally consistent "
        "worlds so decision-makers can spot early signals and avoid betting everything on one outdated model.")

    # 2 METHOD
    h1(doc, "2. Methodology")
    body(doc, "We followed Chermack's three phases:")
    bullet(doc, "Phase 1 (Preparation): purpose statement and system picture.")
    bullet(doc, "Phase 2 (Exploration): STEEP scan, Givens vs Drivers, two peer interviews.")
    bullet(doc, "Phase 3 (Development): 2x2 matrix, four scenario narratives, early-warning indicators, critical reflection.")
    body(doc,
        "A Given is a condition already in motion that every scenario must assume. "
        "A Driver is uncertain in speed or scale and can push futures in different directions.")

    # 3 PHASE 1
    h1(doc, "3. Phase 1: Project Preparation")
    h2(doc, "3.1 Purpose statement")
    body(doc, "By 2035, marketing teams must decide:")
    bullet(doc, "Budget split: SEO, SEM, content, AI visibility / GEO work.")
    bullet(doc, "Team structure: classic SEO vs citation and entity optimization.")
    bullet(doc, "KPIs: website traffic vs in-platform influence and brand mentions in AI answers.")
    bullet(doc, "Content strategy: volume vs authority, proprietary data, and structured markup.")
    body(doc,
        "Our scenarios inform those decisions before a traffic cliff forces them.")

    h2(doc, "3.2 System picture")
    body(doc,
        "Chermack asks you to map the field in layers. Figure 1 shows how we drew it. "
        "The core is simple to say and hard to fix: staying discoverable when AI mediates search.")
    fig(doc, system_img, 5.8, "Figure 1: System picture (four layers from general environment to core issue).")
    bullet(doc, "Core: organic discoverability in AI-mediated search.")
    bullet(doc, "Value chain: users, platforms (Google, ChatGPT, Copilot, Perplexity), publishers and brand sites, agencies, B2B funnels, analytics.")
    bullet(doc, "Specific environment: Germany/EU, SISTRIX and SparkToro data, DMA, AI Act, uneven industry impact.")
    bullet(doc, "General environment: GenAI, copyright, Gartner forecasts, data-centre energy, global platform competition.")

    # 4 STEEP
    h1(doc, "4. Phase 2: STEEP External Analysis")
    body(doc, "Research by Arshi and Ashik. Table 1 is our merged factor list.")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, label in enumerate(["Dimension", "Factor", "Evidence", "Given / Driver", "Source"]):
        table.rows[0].cells[i].text = label
    steep_rows = [
        ("Social", "Zero-click behaviour", "59.7% of EU Google searches end without a click", "Given", "SparkToro 2024"),
        ("Social", "Reliance on AI answers", "AIO on 64.7% of question-based searches", "Driver", "arXiv 2605.14021"),
        ("Technological", "AI search mainstream", "AIO on ~20% of DE queries (~30% long-tail)", "Given", "SISTRIX"),
        ("Technological", "Shift to AI assistants", "Gartner: -25% search volume by 2026", "Driver", "Gartner Feb 2024"),
        ("Technological", "Organic CTR erosion", "CTR on #1: 27% down to 11% when AIO shows", "Driver", "SISTRIX"),
        ("Economic", "Reduced open-web traffic", "374 of 1,000 EU searches reach the open web", "Given", "SparkToro"),
        ("Economic", "Publisher revenue pressure", "~265M lost organic clicks/month in Germany", "Driver", "SISTRIX"),
        ("Economic", "Small publisher decline", "~60% search referral drop over 2 years", "Driver", "Axios"),
        ("Political/Legal", "Mass AI search reach", "AIO reaches 2B+ users globally", "Given", "arXiv 2605.14021"),
        ("Political/Legal", "Regulation uncertainty", "11% of AIO claims unsupported by sources", "Driver", "arXiv 2605.14021"),
        ("Environmental", "AI infrastructure demand", "GenAI search needs far more data-centre energy", "Driver", "IEA, McKinsey, Google"),
    ]
    for row in steep_rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    doc.add_paragraph()
    body(doc, "Table 1: STEEP factors for SEO in the Age of AI (2035 horizon).")

    h2(doc, "4.1 Justifying Givens vs Drivers")
    body(doc,
        "Givens are measured or deployed facts: EU zero-click at 59.7%, AIO live in Germany, "
        "374 clicks per thousand to the open web, 2B+ users on AIO. Every scenario builds on these.")
    body(doc,
        "Drivers are genuine uncertainties: Will AI reliance on question queries keep rising? "
        "Will Gartner's 25% volume drop arrive on schedule? Will CTR pain spread into B2B comparison queries? "
        "Will EU attribution rules tighten after studies found unsupported claims? Will energy costs slow AI search rollout?")
    body(doc,
        "Note: SparkToro also reports higher zero-click in some broader samples (~68% in our earlier reading). "
        "We use the EU figure (59.7%) for European scenario logic and cite the source clearly.")

    h2(doc, "4.2 Industry split (SISTRIX)")
    bullet(doc, "Higher impact so far: how-to, health, recipe, education sites.")
    bullet(doc, "Lower impact so far: e-commerce, brand websites, local businesses.")

    # 5 DATA
    h1(doc, "5. Key Data Insights")
    body(doc, "Figure 2 summarizes the numbers we cite most often in the scenarios.")
    fig(doc, data_img, 6.0, "Figure 2: Key statistics (SparkToro, SISTRIX, Gartner).")
    body(doc,
        "Zero-click moved from about 60.45% in 2024 toward higher EU readings. "
        "When AI Overviews appear, the first organic result loses most of its click power. "
        "Publishers in Germany alone may lose hundreds of millions of organic clicks per month at current rollout levels.")

    # 6 INTERVIEWS
    h1(doc, "6. Stakeholder Interviews")
    body(doc, "Conducted by Ozlem Ceylan, June 2026. Semi-structured, ~10 minutes each, peers outside Team 14.")

    h2(doc, "6.1 Funda Aydin (B2B marketing, GEA Digital)")
    body(doc,
        "Funda has worked 2+ years in marketing at GEA Digital (B2B, industrial engineering). "
        "She told us ChatGPT is her first tool for work research, with a visible shift away from Google over about two years.")
    body(doc,
        'Key quote: "Visibility will become less about ranking number one in Google and more about becoming '
        'a reliable source that AI models repeatedly use."')
    body(doc, "Other insights:")
    bullet(doc, "ChatGPT feels faster because it pulls from many sites and summarizes.")
    bullet(doc, "For routine marketing she does not always check sources; for legal/technical topics she would.")
    bullet(doc, "B2B still needs deep technical content somewhere trustworthy; the click is optional.")
    bullet(doc, "Companies should optimize to be cited and trusted by AI, not only clicked in search results.")
    bullet(doc, "Biggest concern: a few AI platforms becoming gatekeepers of visibility.")

    h2(doc, "6.2 Lorena Contreras (community manager, SEO background)")
    body(doc,
        "Lorena has digital marketing and SEO seminar background; now focused on community enablement. "
        "She still opens Google first when she knows what she wants. She uses ChatGPT to structure text, not to replace search.")
    body(doc,
        'Key quote: "SEO will still exist in 2035, less about keywords, more about credibility and being visible in AI-generated answers."')
    body(doc, "Other insights:")
    bullet(doc, "Google AI summaries can answer without a click even when your content built the answer.")
    bullet(doc, "Trust comes from real examples, expert opinions, customer cases, transparent information.")
    bullet(doc, "Traditional blogs matter less; expert content, video, use cases, and data matter more.")
    bullet(doc, "She would not pay for generic articles AI can reproduce.")

    h2(doc, "6.3 How interviews shaped the matrix")
    body(doc,
        "Funda represents AI-first discovery and citation logic. Lorena represents hybrid discovery and website value for commercial intent. "
        "Both can be true at once for different users and query types. That is why our axes are discovery architecture and value capture, "
        "not SEO dead vs SEO fine.")

    # 7 GIVENS DRIVERS SUMMARY
    h1(doc, "7. Givens and Drivers Summary")
    h2(doc, "7.1 Givens (baseline for all four scenarios)")
    bullet(doc, "59.7% EU zero-click; 374 clicks per 1,000 EU searches (SparkToro).")
    bullet(doc, "AIO on ~20% of DE queries, often above organic #1 (SISTRIX).")
    bullet(doc, "GenAI embedded in retrieval; not experimental.")
    bullet(doc, "AIO global reach 2B+ users.")
    bullet(doc, "Uneven industry impact already visible.")

    h2(doc, "7.2 Drivers (fed into matrix axes)")
    bullet(doc, "Discovery: AI-native default vs hybrid coexistence (Gartner uncertainty, Funda vs Lorena).")
    bullet(doc, "Value capture: zero-click influence vs website funnels (SparkToro, SISTRIX, publisher pain).")

    # 8 MATRIX
    h1(doc, "8. Phase 3: The 2x2 Scenario Matrix")
    body(doc,
        "We selected two critical uncertainties as axes. They are independent enough to produce four distinct futures.")
    h2(doc, "8.1 Axis 1 (vertical): Discovery architecture")
    bullet(doc, "Top: AI-native discovery dominates (ChatGPT, Copilot, answer engines as default).")
    bullet(doc, "Bottom: Hybrid discovery persists (Google + AI + open web coexist).")
    h2(doc, "8.2 Axis 2 (horizontal): Value capture model")
    bullet(doc, "Left: Website-centric funnels still dominate (traffic to lead/sale).")
    bullet(doc, "Right: Zero-click influence economy (citations, mentions, in-platform authority).")

    body(doc, "Figure 3 shows the four quadrants. Table 2 matches scenario names to driver combinations.")
    fig(doc, matrix_img, 6.2, "Figure 3: 2x2 scenario matrix (Horizon 2035).")

    m2 = doc.add_table(rows=1, cols=3)
    m2.style = "Table Grid"
    for i, h in enumerate(["Scenario", "Discovery pole", "Value capture pole"]):
        m2.rows[0].cells[i].text = h
    for row in [
        ("1. Citation Economy", "AI-native", "Zero-click influence"),
        ("2. Last-Click Web", "Hybrid", "Website funnels"),
        ("3. Gatekeeper Web", "AI-native", "Website funnels (expected but broken)"),
        ("4. Slow Adaptation", "Hybrid", "Gradual zero-click shift"),
    ]:
        c = m2.add_row().cells
        for i, v in enumerate(row):
            c[i].text = v
    doc.add_paragraph()
    body(doc, "Table 2: Scenario names and driver combinations.")

    body(doc,
        "Process link (STEEP + interviews to axes): Funda pushed us toward AI-native discovery and citation value. "
        "Lorena anchored hybrid discovery and website value for commercial intent. SparkToro and SISTRIX gave the scale of zero-click pressure.")

    # 9 SCENARIOS
    h1(doc, "9. Four Scenario Narratives (2035)")

    h2(doc, "9.1 Scenario 1: The Citation Economy")
    body(doc, "Drivers: AI-native discovery + zero-click influence economy.")
    body(doc,
        "By 2035 most professional users stopped treating Google as the starting point for work research. "
        "ChatGPT, Copilot, and embedded assistants became the default question interface. Search remained for maps, "
        "logistics, or when you already knew the exact site.")
    body(doc,
        "Marketing fought over citation share: how often your brand, documentation, expert quotes, and structured data "
        "appeared inside AI answers. Agencies rebranded SEO into AI visibility or source authority teams. "
        "B2B industrial firms adapted better than expected. Buyers still needed technical truth, but the win was becoming "
        "the origin layer AI returned to, not chasing a click that might never come.")
    body(doc,
        "Publishers struggled. Some survived via licensing or niche authority. Small sites without leverage folded or moved to newsletters and communities.")
    body(doc, "Early-warning indicators today:")
    bullet(doc, "ChatGPT or Copilot first in B2B research surveys.")
    bullet(doc, "Job postings for GEO / AEO / LLM source optimization.")
    bullet(doc, "EU open-web click share trending below 30%.")
    bullet(doc, "Publisher licensing programs at scale.")

    h2(doc, "9.2 Scenario 2: The Last-Click Web")
    body(doc, "Drivers: Hybrid discovery + website-centric funnels still dominate.")
    body(doc,
        "AI changed search but did not remove the website as the commercial centre. People mixed tools by task. "
        "AIO stayed strong on informational queries; high-intent commercial searches still clicked.")
    body(doc,
        "SEO evolved as Lorena predicted: credibility architecture, schema, expert authorship, customer proof, fast technical sites. "
        "Zero-click hurt, but many B2B teams treated AI answers as top-funnel and invested in onsite experiences summaries could not replace.")
    body(doc, "Early-warning indicators today:")
    bullet(doc, "AIO concentrated in informational queries; commercial CTR relatively stable.")
    bullet(doc, "Pipeline still attributed to organic website sessions.")
    bullet(doc, "Hybrid tool use: same users on Google and AI weekly.")
    bullet(doc, "Structured data spend rising faster than content volume.")

    h2(doc, "9.3 Scenario 3: The Gatekeeper Web")
    body(doc, "Drivers: AI-native discovery + website-centric funnels still expected to work.")
    body(doc,
        "AI platforms became toll roads. Companies kept investing in landing pages because CRM and boards still used funnel metrics. "
        "Users rarely arrived. Traffic flatlined while influence inside AI grew without clean monetization.")
    body(doc,
        "This is Funda's gatekeeper fear at scale. A small set of platforms chose trusted sources. Large brands paid for visibility packages. "
        "Smaller players faced dark influence: mentioned in AI, invisible in analytics. Agencies heard 'we rank but nobody clicks.'")
    body(doc, "Early-warning indicators today:")
    bullet(doc, "Few AI platforms dominate B2B product research journeys.")
    bullet(doc, "Publisher referrals down 50%+ without licensing revenue.")
    bullet(doc, "Paid priority-source or enterprise visibility products.")
    bullet(doc, "Brand mentions in AI absent from web analytics.")

    h2(doc, "9.4 Scenario 4: The Slow Adaptation")
    body(doc, "Drivers: Hybrid discovery + gradual zero-click shift.")
    body(doc,
        "Everyone agreed the old SEO playbook was broken. Almost nobody agreed on the replacement fast enough. "
        "Google mattered. AI mattered. Behaviour varied by age, sector, and task.")
    body(doc,
        "Teams kept hiring for content volume while clicks fell. Generic blogs multiplied. Rankings looked fine while pipeline quality slipped. "
        "Mid-size firms suffered most. Agencies sold rebrands and fear. Survivors experimented early with hybrid measurement and expert content.")
    body(doc, "Early-warning indicators today:")
    bullet(doc, "Flat SEO budgets while AI usage among customers rises.")
    bullet(doc, "Content volume up, conversions and referral traffic down.")
    bullet(doc, "SEO is dead vs SEO is fine camps with no shared metrics.")
    bullet(doc, "Funda-type and Lorena-type behaviours in the same organization.")

    # 10 REFLECTION
    h1(doc, "10. Critical Reflection")
    h2(doc, "10.1 Most plausible scenario")
    body(doc,
        "We lean toward Slow Adaptation. Data shows rising zero-click pressure. Interviews show fragmented behaviour. "
        "SISTRIX's uneven industry impact suggests change is gradual, not a clean flip.")
    h2(doc, "10.2 Most concerning scenario")
    body(doc,
        "Gatekeeper Web. When a few platforms decide visibility, that is a power structure. "
        "Firms measuring websites while users never arrive creates dark influence and unfair dependency.")
    h2(doc, "10.3 Strategic responses by scenario")
    rt = doc.add_table(rows=1, cols=2)
    rt.style = "Table Grid"
    rt.rows[0].cells[0].text = "Scenario"
    rt.rows[0].cells[1].text = "If this world emerges, prioritize"
    for row in [
        ("Citation Economy", "Proprietary data, entities, citation measurement, expert authorship"),
        ("Last-Click Web", "High-intent SEO, technical structure, conversion depth, AI as top-funnel"),
        ("Gatekeeper Web", "Platform diversification, early licensing talks, brand moats AI cannot clone"),
        ("Slow Adaptation", "Parallel experiments, kill generic content, hybrid KPIs before crisis"),
    ]:
        c = rt.add_row().cells
        c[0].text, c[1].text = row

    h2(doc, "10.4 What we would do as future professionals")
    body(doc,
        "Learn discovery strategy, not SEO alone: citations, structured data, authority, analytics beyond sessions. "
        "Push clients to choose which future they prepare for. Take practitioner voices seriously; they often see shifts before dashboards do.")

    # 11 TEAM + REFS
    h1(doc, "11. Team Contributions")
    bullet(doc, "Emran: coordination, purpose, matrix facilitation, moderation, narrative QA, reflection.")
    bullet(doc, "Arshi: STEEP research (Gartner, SparkToro themes), Givens/Drivers tagging.")
    bullet(doc, "Ashik: STEEP merge, source justification, factor table.")
    bullet(doc, "Ozlem: interviews (Funda Aydin, Lorena Contreras), transcription, synthesis.")
    bullet(doc, "Wasay: system picture and scenario visuals, slide deck design.")

    h1(doc, "12. References")
    refs = [
        "SparkToro (2024). Zero-Click Search Study. https://sparktoro.com/blog/2024-zero-click-search-study-for-every-1000-us-google-searches-only-374-clicks-go-to-the-open-web-in-the-eu-its-360/",
        "SISTRIX. AI Overviews in Germany. https://www.sistrix.com/blog/ai-overviews-in-germany/",
        "Gartner (19 Feb 2024). Search Engine Volume Will Drop 25 Percent by 2026. https://www.gartner.com/en/newsroom/press-releases/2024-02-19-gartner-predicts-search-engine-volume-will-drop-25-percent-by-2026-due-to-ai-chatbots-and-other-virtual-agents",
        "arXiv:2605.14021. AI Overview prevalence, claim accuracy, global reach.",
        "Axios. Small publisher search referral decline (~60% over two years).",
        "IEA; McKinsey; Google Sustainability Reports. AI infrastructure energy demand.",
        "Chermack, T. J. Scenario Planning in Organizations.",
        "Team 14 stakeholder interviews: Funda Aydin (17 June 2026); Lorena Contreras (18 June 2026).",
    ]
    for r in refs:
        bullet(doc, r)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
