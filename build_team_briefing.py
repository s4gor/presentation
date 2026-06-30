#!/usr/bin/env python3
"""Team 14 internal briefing pack with 2x2 image and plain-language guide."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_LINE_SPACING
from PIL import Image, ImageDraw, ImageFont

BASE = Path("/Users/emran/Desktop/pptx")
IMG_PATH = BASE / "team_assets" / "matrix_2x2.png"
DOC_PATH = BASE / "Team14_Internal_Briefing.docx"


def make_matrix_image():
    IMG_PATH.parent.mkdir(parents=True, exist_ok=True)
    W, H = 1200, 900
    img = Image.new("RGB", (W, H), "#FAFAF8")
    draw = ImageDraw.Draw(img)

    try:
        title_f = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 28)
        label_f = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 20)
        cell_f = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 17)
        small_f = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 14)
        axis_f = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 16)
    except OSError:
        title_f = label_f = cell_f = small_f = axis_f = ImageFont.load_default()

    draw.text((40, 24), "Team 14 - 2x2 Scenario Matrix (Horizon 2035)", fill="#1a1a1a", font=title_f)

    # Grid area
    gx, gy, gw, gh = 280, 120, 860, 680
    cw, ch = gw // 2, gh // 2

    colors = {
        "gate": ("#FDE8E8", "#C44D6A"),
        "cite": ("#E6F5F4", "#0A7C78"),
        "slow": ("#FEF3C7", "#B8943F"),
        "last": ("#ECFDF5", "#059669"),
    }

    cells = [
        (0, 0, "gate", "3. Gatekeeper Web", "AI runs search.\nSites still matter\non paper.\nUsers don't click.\nPlatforms decide\nwho shows up."),
        (1, 0, "cite", "1. Citation Economy", "AI runs search.\nClicks optional.\nWin by being cited\ninside answers.\nSpecs, data, trust."),
        (0, 1, "slow", "4. Slow Adaptation", "Google + AI both used.\nZero-click creeps up.\nTeams adapt late.\nMessy middle decade."),
        (1, 1, "last", "2. Last-Click Web", "Google + AI both used.\nWebsites still convert.\nSEO evolves.\nB2B still gets clicks\non buying queries."),
    ]

    for col, row, key, title, body in cells:
        x = gx + col * cw
        y = gy + row * ch
        bg, accent = colors[key]
        draw.rectangle([x, y, x + cw - 4, y + ch - 4], fill=bg, outline=accent, width=3)
        draw.text((x + 16, y + 14), title, fill=accent, font=label_f)
        draw.multiline_text((x + 16, y + 52), body, fill="#333333", font=cell_f, spacing=6)

    # Column headers (top)
    draw.text((gx + cw // 2 - 120, 88), "Website funnels still ROI", fill="#555", font=axis_f)
    draw.text((gx + cw + cw // 2 - 110, 88), "Zero-click / citations win", fill="#555", font=axis_f)

    # Row labels (left)
    draw.text((24, gy + ch // 2 - 50), "AI-native\ndiscovery\n(AI first)", fill="#555", font=axis_f)
    draw.text((24, gy + ch + ch // 2 - 50), "Hybrid\ndiscovery\n(Google + AI)", fill="#555", font=axis_f)

    # Axis arrows labels
    draw.text((gx + gw // 2 - 100, H - 42), "VALUE CAPTURE (horizontal axis)", fill="#888", font=small_f)
    draw.text((12, gy + gh // 2 - 12), "DISCOVERY", fill="#888", font=small_f)

    # Legend box
    draw.rectangle([40, H - 88, 520, H - 48], fill="#F0F0EE", outline="#CCC")
    draw.text(
        (52, H - 78),
        "How to read: Funda = up + right tendency. Lorena = down + left tendency. Data = zero-click pushes right.",
        fill="#444",
        font=small_f,
    )

    img.save(IMG_PATH)
    return IMG_PATH


def para(doc, text, size=11, space=10):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(space)
    for r in p.runs:
        r.font.size = Pt(size)
    return p


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def tip(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("Team note: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0x0A, 0x7C, 0x78)
    p.add_run(text)
    p.paragraph_format.space_after = Pt(8)


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    return p


def build_doc(img_path):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    for s in doc.sections:
        s.top_margin = Inches(0.9)
        s.bottom_margin = Inches(0.9)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    h = doc.add_heading("Team 14 Internal Briefing", 0)
    h.alignment = 1
    para(doc, "Not for submission. Read this so everyone knows what we built and why.", size=12)
    para(doc, "Brief B | SEO in the Age of AI | 2035 | Emran, Arshi, Ashik, Ozlem, Wasay")
    doc.add_paragraph()

    heading(doc, "Start here (2-minute version)")
    para(
        doc,
        "Google SEO used to mean: rank high, get clicks, sell on your website. That still happens sometimes. "
        "But a huge share of searches never leave Google. AI Overviews and ChatGPT answer the question before anyone visits you. "
        "Our job in this project was not to guess the one true future. We built four different futures and said what we would do in each.",
    )
    para(
        doc,
        "Two questions drive everything:",
    )
    bullet(doc, "Where do people search in 2035? Mostly AI, or still a mix with Google?")
    bullet(doc, "How do companies win? Website traffic, or being mentioned inside AI answers without a click?")
    tip(doc, "If you only remember one slide, remember the 2x2 matrix below.")

    heading(doc, "The 2x2 matrix (picture + explanation)")
    para(doc, "This is the heart of the project. Vertical = discovery. Horizontal = how money and visibility work.")
    doc.add_picture(str(img_path), width=Inches(6.2))
    para(doc, "Save a copy of the image file here: pptx/team_assets/matrix_2x2.png", size=10)

    heading(doc, "Read the matrix like this", level=2)
    bullet(doc, "TOP row = AI-native discovery. People start in ChatGPT/Copilot, not Google.")
    bullet(doc, "BOTTOM row = Hybrid. Google and AI both stay in daily use (Lorena's world).")
    bullet(doc, "LEFT column = Website funnels still the business model. Traffic matters.")
    bullet(doc, "RIGHT column = Zero-click wins. Being cited in AI answers matters more than clicks.")

    para(doc, "The four boxes:")
    bullet(doc, "Top-left = Gatekeeper Web (bad for small players, Funda's fear)")
    bullet(doc, "Top-right = Citation Economy (optimize to be cited)")
    bullet(doc, "Bottom-left = Slow Adaptation (most realistic messy middle)")
    bullet(doc, "Bottom-right = Last-Click Web (SEO still works, evolved)")

    heading(doc, "What each person contributed")
    bullet(doc, "Emran: keeps the story together, matrix logic, video intro and close")
    bullet(doc, "Arshi: STEEP research (Gartner, SparkToro themes)")
    bullet(doc, "Ashik: merged STEEP table, Given vs Driver tags")
    bullet(doc, "Ozlem: interviews with Funda and Lorena")
    bullet(doc, "Wasay: slides, visuals, system picture")

    heading(doc, "Chermack in normal words (3 phases)")
    bullet(doc, "Phase 1 - Prep: Why scenarios? For whom? Draw the system (who is in the chain).")
    bullet(doc, "Phase 2 - Explore: STEEP scan + interviews. Split facts (Givens) from uncertainties (Drivers).")
    bullet(doc, "Phase 3 - Build: Pick 2 uncertainties as axes. Write 4 stories. Add early warning signs.")

    heading(doc, "STEEP - what we found (short)")
    para(doc, "Arshi and Ashik did the research. Here is what matters for the video, without drowning in numbers.")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, t in enumerate(["Topic", "What it means", "Given or Driver", "One number to remember"]):
        hdr[i].text = t

    rows = [
        ("Zero-click", "Users get answers without visiting sites", "Given", "59.7% EU searches no click"),
        ("AI Overviews", "Google shows AI summary above results", "Given", "~20% of DE searches"),
        ("Open web clicks", "How many searches reach any website", "Given", "374 per 1000 in EU"),
        ("Gartner forecast", "Classic search volume may shrink", "Driver", "down 25% by 2026"),
        ("CTR drop", "Ranking #1 hurts more when AIO shows", "Driver", "27% down to 11%"),
        ("Publisher pain", "Traffic and money under pressure", "Driver", "265M clicks/mo lost DE"),
        ("AI on questions", "Question searches trigger AI a lot", "Driver", "64.7% in one study"),
        ("Bad AI claims", "Regulation might tighten", "Driver", "11% claims unsupported"),
    ]
    for row in rows:
        c = table.add_row().cells
        for i, v in enumerate(row):
            c[i].text = v

    doc.add_paragraph()
    tip(doc, "Given = already happening, all 4 scenarios assume it. Driver = could go different ways, used for matrix.")

    heading(doc, "Interviews - why Funda and Lorena matter")
    heading(doc, "Funda (B2B, GEA Digital)", level=2)
    para(doc, "Opens ChatGPT first, not Google. Thinks B2B still needs deep content, but AI should cite you. Afraid of a few platforms controlling visibility.")
    heading(doc, "Lorena (community, SEO background)", level=2)
    para(doc, "Opens Google first. Uses ChatGPT to write, not to search. Thinks SEO survives but generic blogs are dead. Traffic loss from summaries is real.")
    para(doc, "They are both right for different people. That is literally why we have two axes instead of one SEO good/bad slide.")

    heading(doc, "The four scenarios (what to say in the video)")
    scenarios = [
        (
            "1. Citation Economy (top-right)",
            "Speaker suggestion: Arshi",
            "AI is the front door. Companies win by being the source AI quotes. B2B spec sheets and real data beat blog spam. Publishers struggle unless they license or own a niche.",
        ),
        (
            "2. Last-Click Web (bottom-right)",
            "Speaker suggestion: Ashik",
            "People use Google AND AI. Buying searches still click. SEO becomes trust + structure + expert content. Annoying but workable future.",
        ),
        (
            "3. Gatekeeper Web (top-left)",
            "Speaker suggestion: Ozlem",
            "AI dominates but companies still chase website KPIs. Platforms charge for visibility. You rank but nobody comes. Funda's nightmare version.",
        ),
        (
            "4. Slow Adaptation (bottom-left)",
            "Speaker suggestion: Wasay",
            "Nothing changes cleanly. Zero-click slowly rises. Teams keep old SEO habits too long. Probably the most realistic boring future.",
        ),
    ]
    for title, who, body in scenarios:
        heading(doc, title, level=2)
        para(doc, who, size=10)
        para(doc, body)

    heading(doc, "Reflection (for end of video)")
    bullet(doc, "Most plausible: Slow Adaptation")
    bullet(doc, "Most scary: Gatekeeper Web")
    bullet(doc, "What we'd do: learn discovery strategy, not only classic SEO; measure more than clicks")

    heading(doc, "Video timing cheat sheet")
    timing = [
        ("0:00-1:30", "Emran", "Intro, purpose, Chermack"),
        ("1:30-5:00", "Arshi + Ashik", "STEEP, Givens, Drivers"),
        ("5:00-7:00", "Ozlem", "Interviews"),
        ("7:00-8:00", "Emran", "Show matrix image, explain axes"),
        ("8:00-12:20", "All four scenarios", "About 1 min each"),
        ("12:20-15:00", "Everyone", "Reflection + thanks"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    for i, h in enumerate(["Time", "Who", "What"]):
        t.rows[0].cells[i].text = h
    for row in timing:
        c = t.add_row().cells
        for i, v in enumerate(row):
            c[i].text = v

    doc.add_paragraph()
    heading(doc, "Before we record - checklist")
    bullet(doc, "All 5 faces on screen at least once")
    bullet(doc, "Say out loud: Funda pushed us up/right, Lorena kept hybrid/website")
    bullet(doc, "Do not read slides word for word")
    bullet(doc, "Matrix slide: pause, point at the four boxes")
    bullet(doc, "Swap a sentence or two so it sounds like you")

    heading(doc, "Files in this folder")
    bullet(doc, "Team14_Internal_Briefing.docx (this file)")
    bullet(doc, "team_assets/matrix_2x2.png (matrix image for slides or WhatsApp)")
    bullet(doc, "Horizon_2035_Scenario_Case_Team14.docx (full written case)")
    bullet(doc, "Horizon_2035_Video_Speech_Team14.docx (speaker script)")

    doc.save(DOC_PATH)
    print(DOC_PATH)


if __name__ == "__main__":
    img = make_matrix_image()
    build_doc(img)
