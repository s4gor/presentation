#!/usr/bin/env python3
"""Video speech - natural spoken language for Team 14 recording."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_LINE_SPACING

OUTPUT = "/Users/emran/Desktop/pptx/Horizon_2035_Video_Speech_Team14.docx"


def who(doc, name, when=""):
    p = doc.add_paragraph()
    r = p.add_run(f"{name}" + (f"  [{when}]" if when else ""))
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x0A, 0x7C, 0x78)


def say(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(10)
    for r in p.runs:
        r.font.size = Pt(11)


def slide(doc, text):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    p.paragraph_format.space_after = Pt(6)


def build():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1.1)
        s.right_margin = Inches(1.1)

    h = doc.add_heading("Video Speech - Team 14", 0)
    h.alignment = 1
    say(doc, "Read this like you talk, not like an essay. Pause at the matrix. About 15 minutes total. All five of us speak.")
    doc.add_paragraph()

    # EMRAN
    who(doc, "EMRAN", "0:00 - 1:30")
    slide(doc, "Slide 1 - Title. Everyone on camera if you can.")
    say(doc, "Hey. We're Team 14. I'm Emran. With me are Arshi, Ashik, Ozlem, and Wasay.")
    say(doc, "Our topic is Brief B: SEO in the Age of AI. Horizon 2035.")
    say(doc, "Quick thing up front. We're not here to predict the future. Our professor didn't ask for that. Chermack's whole method is: the future is uncertain, so build a few different versions of it and ask what you'd do in each one. That's what we did.")
    slide(doc, "Slide 3 - Purpose")
    say(doc, "Who cares about this? Brand people. Agencies. Publishers. B2B marketers. Anyone who still has to explain a budget when half the room thinks SEO is dead and the other half thinks nothing changed.")
    say(doc, "By 2035 they'll need to decide: do we still pay for classic SEO? Do we optimize to show up inside ChatGPT answers? Do we even measure clicks anymore?")
    say(doc, "That's phase one in Chermack. Preparation. Purpose. Wasay drew the system picture on the next slides. I'll keep it short: the core problem is getting found when AI sits between you and the customer.")
    say(doc, "Arshi and Ashik did the research scan. Ozlem talked to real people. I'll come back for the matrix and the wrap-up.")

    # ARSHI
    who(doc, "ARSHI", "1:30 - 3:00")
    slide(doc, "Slide 6 - STEEP")
    say(doc, "Hey, Arshi. So Ashik and I both did STEEP research and honestly there was overlap. We merged it into one table so we're not saying the same SparkToro number four times.")
    say(doc, "Social side: zero-click is already normal. In the EU, SparkToro says about sixty percent of Google searches end with no click to any website. Out of a thousand searches, only three hundred seventy-four send someone to the open web. That's not a prediction. That's what's happening now.")
    say(doc, "Tech side: SISTRIX looked at over a hundred million keywords in Germany. AI Overviews show up on roughly twenty percent of searches. Long-tail, it's closer to thirty. And most of the time the AI box is above the normal number one result.")
    say(doc, "Economics: it hits publishers and how-to sites harder than e-commerce or local business. Political: AI Overviews reach billions of people, so regulators are watching. Environment: AI search burns more energy than old-school search. That might slow things down later. Might not.")
    say(doc, "Ashik will explain what we tagged as Given versus Driver. That's the part that actually builds the matrix.")

    # ASHIK
    who(doc, "ASHIK", "3:00 - 5:00")
    slide(doc, "Slides 7-8 - Givens, Drivers, Data")
    say(doc, "Ashik here. Givens versus Drivers. Simple version.")
    say(doc, "A Given is something already true enough that every future story has to start with it. Zero-click behaviour. AI Overviews live in Google Germany. Three hundred seventy-four clicks per thousand in the EU. Two billion people can see AI Overviews. You don't scenario-plan whether those exist. You plan inside them.")
    say(doc, "A Driver is something that could go different ways. Gartner says search volume might drop twenty-five percent by twenty-twenty-six. Maybe faster, maybe slower. When an Overview shows, click-through on number one crashes from twenty-seven to eleven percent. Will that spread to every industry? Unclear. Small publishers already down a lot on search traffic. Axios wrote about that.")
    say(doc, "We fought a bit on the tags. Good. If everything's a Given, you have no matrix. If everything's a Driver, it's noise.")
    say(doc, "One number slide worth remembering: sixty-four percent of question-type searches trigger AI Overviews in one big study. And eleven percent of AI claims weren't fully supported by the sources cited. Regulation might get stricter because of stuff like that.")
    say(doc, "Ozlem's interviews matter here because stats don't tell you how people actually behave at work.")

    # OZLEM
    who(doc, "OZLEM", "5:00 - 7:00")
    slide(doc, "Slides 9-10 - Funda and Lorena")
    say(doc, "Ozlem. I did two interviews, about ten minutes each, with people from other teams.")
    say(doc, "First, Funda Aydin. She works in B2B marketing at GEA Digital. I asked what she opens first when she needs info for work. She said ChatGPT. Not Google. She told me that's changed over the last two years.")
    say(doc, "She said something we put on the slide: visibility is less about ranking number one and more about being a source AI keeps using. She still thinks B2B needs serious technical content somewhere. But that somewhere might not need a click.")
    say(doc, "And her worry? A few AI platforms deciding who gets seen. She said that plainly. We remembered it.")
    say(doc, "Second person, Lorena Contreras. Community manager, some SEO background. Different story. Google is still her first move when she knows what she wants. ChatGPT helps her write and organize. Not replace search.")
    say(doc, "She's seen Google put the answer on top and the company that wrote the content gets no visit. She thinks SEO is still there in twenty-thirty-five but it's credibility, structure, expert stuff. Not keyword blogs. She wouldn't pay for generic articles AI can copy.")
    say(doc, "So Funda and Lorena aren't debating who's right. They're both right for different people and different tasks. Emran's gonna show how that became our matrix.")

    # EMRAN MATRIX
    who(doc, "EMRAN", "7:00 - 8:15")
    slide(doc, "Slide 12 - 2x2 matrix image. SLOW DOWN HERE.")
    say(doc, "Back to me. Phase three. Development. This is the slide you don't rush.")
    say(doc, "Two questions. One: where do people search in twenty-thirty-five? Mostly AI? Or still a mix with Google? That's the vertical axis. Top row, AI-first. Bottom row, hybrid, both still in daily use.")
    say(doc, "Two: how do companies win? Website traffic and funnels? Or being mentioned and cited inside AI answers without a click? That's horizontal. Left, websites still matter for money. Right, zero-click influence.")
    say(doc, "Four boxes. Top right: Citation Economy. Bottom right: Last-Click Web. Top left: Gatekeeper Web. Bottom left: Slow Adaptation.")
    say(doc, "Say this out loud because Chermack wants to hear the logic: Funda pushed us up and to the right. Lorena kept us on hybrid and website value. SparkToro and SISTRIX gave us the scale on zero-click. That's how STEEP plus interviews became axes. Not SEO good versus SEO bad. That would've been lazy.")

    # SCENARIO 1 ARSHI
    who(doc, "ARSHI", "8:15 - 9:15")
    slide(doc, "Slide 13 - Citation Economy")
    say(doc, "Scenario one. Citation Economy. Top right of the matrix.")
    say(doc, "Picture twenty-thirty-five. You start research in ChatGPT or Copilot, not Google. Marketing teams stop fighting for blue link number one. They fight to get cited in the answer. Who owns the spec sheet AI keeps quoting? Who owns the case study?")
    say(doc, "B2B actually does okay here if you publish real technical depth. Thin blog spam dies. Some publishers survive through licensing or being too niche to ignore.")
    say(doc, "Warning signs you'd see early: job posts saying GEO instead of SEO. More people saying ChatGPT is step one at work. EU click numbers keep sliding.")

    # SCENARIO 2 ASHIK
    who(doc, "ASHIK", "9:15 - 10:15")
    slide(doc, "Slide 14 - Last-Click Web")
    say(doc, "Scenario two. Last-Click Web. Bottom right.")
    say(doc, "AI changed things but the website didn't die for sales. People use Google and AI for different jobs. Ask ChatGPT for an overview, Google a vendor, click through when you're actually buying.")
    say(doc, "SEO looks more like what Lorena said. Trust, structure, proof, fast sites. Not stuffing keywords. Annoying future maybe, but workable. A lot of B2B still gets pipeline from organic if the query is commercial enough.")

    # SCENARIO 3 OZLEM
    who(doc, "OZLEM", "10:15 - 11:15")
    slide(doc, "Slide 15 - Gatekeeper Web")
    say(doc, "Scenario three. Gatekeeper Web. Top left. This one's the one that scares me.")
    say(doc, "AI runs discovery but companies still build everything around website funnels because that's how CRM and boards think. Users don't come. You rank, sort of, but invisibly. Platforms charge for visibility. Funda warned us about gatekeepers. This is that world fully grown up.")
    say(doc, "Small publishers and smaller B2B brands get squeezed. You show up in an answer maybe, but analytics say nothing. Dark influence.")

    # SCENARIO 4 WASAY
    who(doc, "WASAY", "11:15 - 12:15")
    slide(doc, "Slide 16 - Slow Adaptation")
    say(doc, "Wasay. Scenario four. Slow Adaptation. Bottom left.")
    say(doc, "Honestly this feels like the most realistic boring future. Everyone knows the old SEO playbook is broken. Nobody replaces it cleanly. Zero-click creeps up. Teams keep hiring content volume. Agencies rebrand SEO and sell fear.")
    say(doc, "Google still matters. AI still matters. Funda and Lorena both exist in the same company and confuse the strategy deck. Survivors are the ones who experimented early, even messily.")

    # REFLECTION
    who(doc, "EMRAN", "12:15 - 13:30")
    slide(doc, "Slide 18 - Reflection")
    say(doc, "Chermack asks three things at the end. We tried to answer them straight.")
    say(doc, "Most plausible? Slow Adaptation. Data pressure is real. People change habits slowly. Big companies don't pivot clean.")
    say(doc, "Most worrying? Gatekeeper Web. Power concentrated in a few platforms is not a tech detail. It's a business risk.")
    say(doc, "What would we do? Stop treating SEO like one skill. Learn discovery strategy. Citations, structure, authority, metrics that aren't only clicks.")
    say(doc, "Quick round. One line each.")

    who(doc, "ARSHI", "13:30")
    say(doc, "I'd tell a client to pick which box they're betting on before the traffic cliff, not after.")

    who(doc, "ASHIK", "13:40")
    say(doc, "I'd cut generic content spend and put money into expert content and technical setup.")

    who(doc, "OZLEM", "13:50")
    say(doc, "I'd keep talking to people like Funda and Lorena every year. Practitioners see it before the dashboard.")

    who(doc, "WASAY", "14:00")
    say(doc, "I'd measure AI visibility somehow, not only Search Console.")

    who(doc, "EMRAN", "14:10 - 15:00")
    slide(doc, "Slide 19 - Thank you / sources")
    say(doc, "Sources on the slide: SparkToro, SISTRIX, Gartner, Chermack, our two interviews.")
    say(doc, "We prepared with Chermack. We explored with STEEP and stakeholder voices. We built four scenarios with early warnings. They're tools for thinking, not fortune telling.")
    say(doc, "Thanks for watching. Happy to take questions.")

    doc.add_paragraph()
    say(doc, "Recording tips: don't read word-for-word if it sounds stiff. Look at the camera on the matrix slide. All five faces visible at least once. Total time 14-16 minutes is fine.")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
